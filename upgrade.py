# ============================================================
# 🔥 CRITICAL RUNTIME FIX (DO NOT REMOVE)
# Fixes:
# 1. Streamlit + PyTorch torch.classes crash
# 2. Python 3.13 asyncio "no running event loop"
# ============================================================

import os
os.environ["STREAMLIT_WATCHER_TYPE"] = "none"
os.environ["TOKENIZERS_PARALLELISM"] = "false"

import asyncio
try:
    asyncio.get_running_loop()
except RuntimeError:
    asyncio.set_event_loop(asyncio.new_event_loop())

import sys
import types

fake_torch_classes = types.SimpleNamespace()
sys.modules["torch.classes"] = fake_torch_classes

# ============================================================
# UNIVERSAL RAG – Diagram Aligned, No Hallucination
# Supports: CSV | PDF | TXT | HTML | DOCX
# Retrieval: BM25 (CSV) + Chroma Vector (Docs)
# UI: Streamlit
# LLM: ChatOllama (qwen2.5:7b)
# ============================================================

import re
import io
import time
import tempfile
import pandas as pd
import streamlit as st
from io import BytesIO
from collections import deque
from datetime import datetime

from rank_bm25 import BM25Okapi
from docx import Document as WordDoc

from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_ollama import ChatOllama
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser

from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    BSHTMLLoader,
    Docx2txtLoader
)

# ============================================================
# UTILITIES
# ============================================================

def normalize_id(text: str) -> str:
    return re.sub(r'[^a-zA-Z0-9]', '', text).lower() if text else ""

def extract_req_id(text: str):
    m = re.search(r"req[\s\-_]?\d+", text, re.IGNORECASE)
    return normalize_id(m.group(0)) if m else None

def is_csv_intent(query: str) -> bool:
    keywords = ["row", "status", "id", "requirement", "category", "table", "csv", "value"]
    q = query.lower()
    return extract_req_id(q) is not None or any(k in q for k in keywords)

def find_in_history(history, query):
    q = query.strip().lower()
    for h in reversed(history):
        if h["question"].strip().lower() == q:
            return h["answer"], h["sources"]
    return None, None

# ============================================================
# HISTORY EXPORT
# ============================================================

def export_history_csv(history):
    df = pd.DataFrame(history)
    buffer = io.StringIO()
    df.to_csv(buffer, index=False)
    return buffer.getvalue()

def export_history_word(history):
    doc = WordDoc()
    doc.add_heading("DocIntel AI – Chat History", level=1)

    for i, h in enumerate(history, 1):
        doc.add_paragraph(f"Q{i}: {h['question']}", style="Heading 2")
        doc.add_paragraph(f"Time: {h['timestamp']}")
        doc.add_paragraph(h["answer"])
        if h["sources"]:
            p = doc.add_paragraph()
            p.add_run(f"Sources: {', '.join(h['sources'])}").italic = True

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ============================================================
# UNIVERSAL RAG CORE (BASE DOCUMENTS PRESERVED)
# ============================================================

class UniversalRAG:
    def __init__(self):
        self.embeddings = HuggingFaceEmbeddings(
            model_name="all-MiniLM-L6-v2",
            model_kwargs={"device": "cpu"},
            encode_kwargs={"normalize_embeddings": True}
        )

        self.llm = ChatOllama(
            model="qwen2.5:7b",
            temperature=0
        )

        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=150
        )

        self.vector_db = None
        self.bm25_indexes = {}
        self.bm25_docs = {}

    # ---------------- CSV (UNCHANGED) ----------------
    def load_csv(self, path):
        fname = os.path.basename(path)
        df = pd.read_csv(path).fillna("N/A")

        docs, corpus = [], []

        for idx, row in df.iterrows():
            raw = " | ".join(f"{k}: {v}" for k, v in row.items())
            norm_id = extract_req_id(raw)

            doc = Document(
                page_content=f"[Source: {fname}]\n{raw}",
                metadata={
                    "source": fname,
                    "row": idx + 1,
                    "type": "csv",
                    "norm_id": norm_id
                }
            )

            docs.append(doc)
            corpus.append(re.findall(r"\w+", raw.lower()))

        self.bm25_indexes[fname] = BM25Okapi(corpus)
        self.bm25_docs[fname] = docs

    # ---------------- FILE ROUTER ----------------
    def process_file(self, path):
        name = path.lower()

        if name.endswith(".csv"):
            self.load_csv(path)
            return []

        if name.endswith(".pdf"):
            return PyPDFLoader(path).load_and_split(self.splitter)
        if name.endswith(".docx"):
            return Docx2txtLoader(path).load_and_split(self.splitter)
        if name.endswith(".txt"):
            return TextLoader(path).load_and_split(self.splitter)
        if name.endswith(".html"):
            return BSHTMLLoader(path).load_and_split(self.splitter)

        return []

    # ---------------- BASE DOCUMENTS PRESERVED ----------------
    def load_documents(self, folder):
        docs = []
        for f in os.listdir(folder):
            docs.extend(self.process_file(os.path.join(folder, f)))

        if docs:
            if self.vector_db:
                self.vector_db.add_documents(docs)
            else:
                self.vector_db = Chroma.from_documents(docs, self.embeddings)

    # ---------------- SEARCH ----------------
    def search(self, query):
        csv_results, doc_results = [], []
        q_id = extract_req_id(query)
        tokens = re.findall(r"\w+", query.lower())

        # 🔥 ALWAYS search CSV (not only csv intent)
        for fname, bm25 in self.bm25_indexes.items():
            scores = bm25.get_scores(tokens)
        for idx, score in enumerate(scores):
            doc = self.bm25_docs[fname][idx]

            # Strong boost for explicit req ID
            if q_id and doc.metadata["norm_id"] == q_id:
                score += 1000

            # Medium boost if requirement text matches
            if "requirement_text" in doc.page_content.lower():
                score += 5

            if score > 0:
                csv_results.append((doc, score))

        csv_results = [
            d for d, _ in sorted(csv_results, key=lambda x: x[1], reverse=True)[:5]
        ]

        # Vector DB (PDF / DOCX)
        if self.vector_db:
            doc_results = self.vector_db.similarity_search(query, k=4)

        return csv_results, doc_results

    # ---------------- ANSWER ----------------
    def answer(self, query):
        csv_docs, doc_docs = self.search(query)
        all_docs = csv_docs + doc_docs

        if not all_docs:
            return "No matching information found.", []

        context = "\n\n---\n\n".join(d.page_content for d in all_docs)

        prompt = ChatPromptTemplate.from_template(
            "Use ONLY the context below.\n\nContext:\n{context}\n\nQuestion:\n{question}\n\nAnswer:"
        )

        answer = (
            prompt | self.llm | StrOutputParser()
        ).invoke({"context": context, "question": query})

        sources = list(dict.fromkeys(
            d.metadata.get("source", "Unknown") for d in all_docs
        ))[:2]

        return answer, sources

# ============================================================
# STREAMLIT UI
# ============================================================

st.set_page_config("DocIntel AI", layout="wide", page_icon="🔍")
st.title("🔍 DocIntel AI – Universal RAG")

if "history" not in st.session_state:
    st.session_state.history = deque(maxlen=50)

if "temp_dir" not in st.session_state:
    st.session_state.temp_dir = tempfile.mkdtemp()

if "rag" not in st.session_state:
    st.session_state.rag = UniversalRAG()

    # ✅ BASE DOCUMENTS AUTO-EMBED (UNCHANGED BEHAVIOR)
    if os.path.exists("./base_documents"):
        st.session_state.rag.load_documents("./base_documents")

with st.sidebar:
    st.header("📂 Upload Documents")
    uploads = st.file_uploader(
        "Upload Files",
        type=["csv", "pdf", "docx", "txt", "html"],
        accept_multiple_files=True
    )

    if uploads and st.button("Process & Index"):
        for f in uploads:
            with open(os.path.join(st.session_state.temp_dir, f.name), "wb") as out:
                out.write(f.read())
        st.session_state.rag.load_documents(st.session_state.temp_dir)
        st.success("Indexing completed!")

    if st.session_state.history:
        st.download_button(
            "📄 CSV",
            export_history_csv(list(st.session_state.history)),
            "chat_history.csv"
        )
        st.download_button(
            "📝 Word",
            export_history_word(list(st.session_state.history)),
            "chat_history.docx"
        )

for h in st.session_state.history:
    with st.chat_message("user"):
        st.markdown(h["question"])
    with st.chat_message("assistant"):
        st.markdown(h["answer"])
        if h["sources"]:
            st.caption(f"Sources: {', '.join(h['sources'])}")

if query := st.chat_input("Ask about your files..."):
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    cached_answer, cached_sources = find_in_history(
        st.session_state.history, query
    )

    with st.chat_message("assistant"):
        if cached_answer:
            st.markdown(cached_answer)
            st.caption(f"Sources: {', '.join(cached_sources)}")
        else:
            with st.spinner("Analyzing..."):
                ans, src = st.session_state.rag.answer(query)
                st.markdown(ans)
                if src:
                    st.caption(f"Sources: {', '.join(src)}")

    st.session_state.history.append({
        "timestamp": ts,
        "question": query,
        "answer": cached_answer or ans,
        "sources": cached_sources or src
    })
